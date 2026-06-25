import csv
import json
import zipfile
import base64
import warnings
from pathlib import Path
from collections import Counter, defaultdict

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

warnings.filterwarnings("ignore")
pd.set_option("display.max_columns", 200)
pd.set_option("display.max_rows", 200)
sns.set_theme(style="whitegrid", palette="Set2")


# ============================================================
# CONFIG
# ============================================================

DATA_PATH = Path(r"D:\Code\DA\1")

# If running locally, use:
# DATA_PATH = Path(r"C:\Users\sangk\Downloads")

OUT_PATH = DATA_PATH / "step3_outputs"
FIG_PATH = OUT_PATH / "figures"
TABLE_PATH = OUT_PATH / "tables"

FIG_PATH.mkdir(parents=True, exist_ok=True)
TABLE_PATH.mkdir(parents=True, exist_ok=True)

FILES = {
    "application_train": "application_train.csv",
    "application_test": "application_test.csv",
    "bureau": "bureau.csv",
    "bureau_balance": "bureau_balance.csv",
    "previous_application": "previous_application.csv",
    "POS_CASH_balance": "POS_CASH_balance.csv",
    "installments_payments": "installments_payments.csv",
    "credit_card_balance": "credit_card_balance.csv",
    "sample_submission": "sample_submission.csv",
    "columns_description": "HomeCredit_columns_description.csv",
}

KEY_SPECS = {
    "application_train.csv": ["SK_ID_CURR"],
    "application_test.csv": ["SK_ID_CURR"],
    "bureau.csv": ["SK_ID_CURR", "SK_ID_BUREAU"],
    "bureau_balance.csv": ["SK_ID_BUREAU"],
    "previous_application.csv": ["SK_ID_CURR", "SK_ID_PREV"],
    "POS_CASH_balance.csv": ["SK_ID_CURR", "SK_ID_PREV"],
    "installments_payments.csv": ["SK_ID_CURR", "SK_ID_PREV"],
    "credit_card_balance.csv": ["SK_ID_CURR", "SK_ID_PREV"],
    "sample_submission.csv": ["SK_ID_CURR"],
}


# ============================================================
# HELPERS
# ============================================================

def get_encoding(file_name):
    return "cp1252" if file_name == "HomeCredit_columns_description.csv" else "utf-8"


def read_header(file_name):
    path = DATA_PATH / file_name
    with open(path, "r", encoding=get_encoding(file_name), newline="") as f:
        return next(csv.reader(f))


def save_table(df, name):
    path = TABLE_PATH / f"{name}.csv"
    df.to_csv(path, index=False)
    print(f"Saved table: {path}")
    print(df.head(30).to_string(index=False))


def save_fig(name):
    path = FIG_PATH / f"{name}.png"
    plt.tight_layout()
    plt.savefig(path, dpi=160, bbox_inches="tight")
    plt.show()
    print(f"Saved figure: {path}")


def format_value_for_report(value):
    if pd.isna(value):
        return ""
    if isinstance(value, (float, np.floating)):
        return f"{value:,.4f}"
    if isinstance(value, (int, np.integer)):
        return f"{value:,}"
    return str(value)


def read_report_table(file_name, max_rows=None):
    path = TABLE_PATH / file_name
    if not path.exists():
        return None
    df = pd.read_csv(path)
    if max_rows is not None:
        return df.head(max_rows)
    return df


def add_docx_dataframe(document, df, max_rows=None):
    if max_rows is not None:
        df = df.head(max_rows)
    if df is None or df.empty:
        document.add_paragraph("No rows available.")
        return
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = format_value_for_report(row[col])


def dataframe_to_html_table(df, max_rows=None):
    if df is None or df.empty:
        return "<p>No rows available.</p>"
    if max_rows is not None:
        df = df.head(max_rows)
    return df.to_html(index=False, escape=True, border=0)


def image_to_base64_html(image_path):
    if image_path is None or not image_path.exists():
        return ""
    encoded = base64.b64encode(image_path.read_bytes()).decode("utf-8")
    return f'<img src="data:image/png;base64,{encoded}" />'


def build_step3_word_report():
    report_sections = [
        {
            "title": "1. Key Quality Profile",
            "why": "Kiểm tra null key và độ lặp của các key chính để đảm bảo không merge hoặc đếm sai cấp độ dữ liệu.",
            "table": "01_key_quality_profile.csv",
            "figure": "01_null_values_in_key_columns.png",
            "max_rows": None,
        },
        {
            "title": "2. Missing Summary By Table",
            "why": "Tóm tắt mức độ thiếu dữ liệu theo từng bảng để xác định bảng/cột cần ưu tiên ở bước cleaning.",
            "table": "03_missing_summary_by_table.csv",
            "figure": "02_max_missing_rate_by_table.png",
            "max_rows": None,
        },
        {
            "title": "3. Top Missing Columns Overall",
            "why": "Xác định các cột có missing cao nhất, đặc biệt là nhóm biến nhà ở, tài sản và các cột có khả năng mang ý nghĩa business.",
            "table": "04_top30_missing_columns_overall.csv",
            "figure": "03_top30_missing_columns_overall.png",
            "max_rows": 30,
        },
        {
            "title": "4. Missingness vs TARGET",
            "why": "Đánh giá missing có phải tín hiệu liên quan đến default hay không bằng cách so sánh default rate giữa nhóm missing và non-missing.",
            "table": "05_application_missingness_vs_target.csv",
            "figure": "04_missingness_vs_target_default_rate_diff.png",
            "max_rows": 40,
        },
        {
            "title": "5. Special Values",
            "why": "Ghi nhận các giá trị đặc biệt như DAYS_EMPLOYED = 365243 hoặc CODE_GENDER = XNA để xử lý có kiểm soát thay vì xóa tùy tiện.",
            "table": "06_application_special_values.csv",
            "figure": "05_application_special_values.png",
            "max_rows": None,
        },
        {
            "title": "6. Special Values vs TARGET",
            "why": "Kiểm tra các special values có liên quan đến default rate hay không trước khi quyết định replace, impute hoặc thêm indicator.",
            "table": "07_special_values_vs_target.csv",
            "figure": None,
            "max_rows": None,
        },
        {
            "title": "7. Time-like Variable Quality",
            "why": "Kiểm tra biến ngày/tháng âm, giá trị dương bất thường và special codes để tránh diễn giải sai thời gian.",
            "table": "08_time_like_variable_quality.csv",
            "figure": "06_positive_time_like_values.png",
            "max_rows": None,
        },
        {
            "title": "8. Plausibility Checks",
            "why": "Kiểm tra các điều kiện hợp lý về business như amount âm, debt lớn hơn credit, trả chậm, trả thiếu và DPD bất thường.",
            "table": "09_plausibility_checks.csv",
            "figure": "07_top_plausibility_checks.png",
            "max_rows": None,
        },
        {
            "title": "9. Application Core Outlier Profile",
            "why": "Dùng percentile để hiểu phân phối biến lõi như income, credit, annuity và các affordability ratios.",
            "table": "10_application_core_outlier_profile.csv",
            "figure": "08_application_outlier_visual_check.png",
            "max_rows": None,
        },
        {
            "title": "10. Train-Test Drift PSI",
            "why": "Đánh giá train và test có phân phối tương đồng không. PSI cao là tín hiệu cần xem kỹ trước khi model hóa.",
            "table": "11_train_test_drift_psi.csv",
            "figure": "09_train_test_drift_psi.png",
            "max_rows": None,
        },
        {
            "title": "11. Cleaning Rules",
            "why": "Tổng hợp các quyết định cleaning chính, lý do, alternative và pros/cons để bảo vệ tính đúng đắn của project.",
            "table": "12_cleaning_rules.csv",
            "figure": None,
            "max_rows": None,
        },
        {
            "title": "12. Cleaning Impact Summary",
            "why": "Cho biết mỗi rule ảnh hưởng bao nhiêu dòng, giúp đánh giá mức độ tác động trước khi áp dụng chính thức.",
            "table": "14_cleaning_impact_summary.csv",
            "figure": None,
            "max_rows": None,
        },
    ]

    docx_path = OUT_PATH / "step3_data_quality_report.docx"
    html_path = OUT_PATH / "step3_data_quality_report.html"

    def write_html_report():
        styles = """
        <style>
        body { font-family: Arial, sans-serif; margin: 32px; color: #1f2937; }
        h1 { color: #111827; border-bottom: 2px solid #e5e7eb; padding-bottom: 6px; }
        h2 { color: #374151; margin-top: 28px; }
        table { border-collapse: collapse; width: 100%; margin: 12px 0 24px 0; font-size: 12px; }
        th, td { border: 1px solid #d1d5db; padding: 6px; text-align: left; vertical-align: top; }
        th { background: #f3f4f6; }
        img { max-width: 100%; margin: 12px 0 24px 0; border: 1px solid #e5e7eb; }
        .note { background: #f9fafb; border-left: 4px solid #60a5fa; padding: 12px; }
        </style>
        """
        parts = [
            "<html><head><meta charset='utf-8'>",
            styles,
            "</head><body>",
            "<h1>Step 3 - Data Quality Assessment & Cleaning Strategy</h1>",
            "<p class='note'>Report này tổng hợp kết quả kiểm tra chất lượng dữ liệu. File HTML được tạo lại ở mỗi lần chạy để đồng bộ với bảng CSV và Word report.</p>",
            "<h2>Executive Summary</h2>",
            "<ul>",
            "<li>Step 3 phân biệt lỗi kỹ thuật, missing có ý nghĩa business và special values cần giữ tín hiệu.</li>",
            "<li>Aggregation chính thức sẽ được thực hiện ở Step 4.</li>",
            "<li>Các quyết định cleaning quan trọng gồm DAYS_EMPLOYED = 365243, previous_application date special values, CODE_GENDER = XNA, property missingness, credit card zero limit, outliers và train-test drift.</li>",
            "</ul>",
        ]

        for section in report_sections:
            parts.append(f"<h2>{section['title']}</h2>")
            parts.append(f"<p>{section['why']}</p>")
            table_df = read_report_table(section["table"], max_rows=section["max_rows"])
            parts.append(dataframe_to_html_table(table_df, max_rows=section["max_rows"]))
            if section["figure"]:
                parts.append(image_to_base64_html(FIG_PATH / section["figure"]))

        parts.append("<h2>Final Note</h2>")
        parts.append("<p>Các bảng CSV và hình PNG đầy đủ vẫn được lưu trong thư mục step3_outputs.</p>")
        parts.append("</body></html>")
        html_path.write_text("\n".join(parts), encoding="utf-8")
        print(f"Saved HTML report: {html_path}")
        return html_path

    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        document = Document()
        document.add_heading("Step 3 - Data Quality Assessment & Cleaning Strategy", level=0)
        intro = (
            "Report này tổng hợp kết quả kiểm tra chất lượng dữ liệu, bao gồm key quality, missing values, "
            "special values, time-like variables, plausibility checks, outlier profile, train-test drift và cleaning rules. "
            "Mục tiêu là xác định dữ liệu có vấn đề gì, vì sao vấn đề đó quan trọng, và quy tắc cleaning nào nên được dùng trước khi tạo bảng phân tích cuối cùng."
        )
        p = document.add_paragraph(intro)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        document.add_heading("Executive Summary", level=1)
        for bullet in [
            "Step 3 không chỉ tìm lỗi dữ liệu mà còn phân biệt lỗi kỹ thuật, missing có ý nghĩa business và special values cần giữ tín hiệu.",
            "Các bảng lịch sử vẫn chưa được aggregate ở bước này; aggregation chính thức sẽ được thực hiện ở Step 4.",
            "Các quyết định cleaning quan trọng nhất gồm xử lý DAYS_EMPLOYED = 365243, CODE_GENDER = XNA, missing values, monetary outliers và train-test drift.",
            "Các biểu đồ trong report phục vụ đánh giá chất lượng dữ liệu, chưa phải dashboard cuối cùng.",
        ]:
            document.add_paragraph(bullet, style="List Bullet")

        for section in report_sections:
            document.add_heading(section["title"], level=1)
            document.add_paragraph(section["why"])

            table_df = read_report_table(section["table"], max_rows=section["max_rows"])
            if table_df is not None:
                document.add_heading("Result Table", level=2)
                add_docx_dataframe(document, table_df, max_rows=section["max_rows"])

            if section["figure"]:
                fig_path = FIG_PATH / section["figure"]
                if fig_path.exists():
                    document.add_heading("Figure", level=2)
                    document.add_picture(str(fig_path), width=Inches(6.5))

        document.add_heading("Final Note", level=1)
        document.add_paragraph(
            "File Word này là report tổng hợp. Các bảng CSV và hình PNG đầy đủ vẫn được lưu trong thư mục step3_outputs "
            "để có thể kiểm tra hoặc đưa vào dashboard/report riêng."
        )
        document.save(docx_path)
        print(f"Saved Word report: {docx_path}")
        write_html_report()
        return docx_path

    except Exception as exc:
        print(f"Could not create .docx report because: {exc}")
        print("Creating HTML report instead. Microsoft Word can open this HTML file.")
        return write_html_report()


def read_csv_chunks(file_name, usecols=None, chunksize=500_000, low_memory=False):
    return pd.read_csv(
        DATA_PATH / file_name,
        usecols=usecols,
        chunksize=chunksize,
        encoding=get_encoding(file_name),
        low_memory=low_memory,
    )


def safe_divide(a, b):
    if b == 0 or pd.isna(b):
        return np.nan
    return a / b


def exact_missing_profile(file_name, chunksize=500_000):
    total_rows = 0
    missing_counts = None
    dtypes_seen = {}

    for chunk in read_csv_chunks(file_name, chunksize=chunksize, low_memory=False):
        total_rows += len(chunk)
        chunk_missing = chunk.isna().sum()
        missing_counts = chunk_missing if missing_counts is None else missing_counts.add(chunk_missing, fill_value=0)
        for col, dtype in chunk.dtypes.items():
            dtypes_seen.setdefault(col, str(dtype))

    result = missing_counts.reset_index()
    result.columns = ["column", "missing_count"]
    result["missing_pct"] = result["missing_count"] / total_rows * 100
    result["file"] = file_name
    result["rows"] = total_rows
    result["dtype_first_chunk"] = result["column"].map(dtypes_seen)
    return result[["file", "column", "dtype_first_chunk", "rows", "missing_count", "missing_pct"]].sort_values(
        ["missing_pct", "missing_count"], ascending=False
    )


def key_quality_profile(file_name, key_cols, chunksize=500_000):
    total_rows = 0
    null_counts = {col: 0 for col in key_cols}
    unique_sets = {col: set() for col in key_cols}

    for chunk in read_csv_chunks(file_name, usecols=key_cols, chunksize=chunksize):
        total_rows += len(chunk)
        for col in key_cols:
            null_counts[col] += int(chunk[col].isna().sum())
            unique_sets[col].update(chunk[col].dropna().astype(str).unique())

    rows = []
    for col in key_cols:
        unique_count = len(unique_sets[col])
        rows.append({
            "file": file_name,
            "key": col,
            "rows": total_rows,
            "null_key_count": null_counts[col],
            "unique_values": unique_count,
            "duplicate_rows_vs_single_key": total_rows - unique_count,
            "null_key_pct": null_counts[col] / total_rows * 100,
        })
    return rows


def numeric_percentile_profile(df, columns, percentiles=None):
    if percentiles is None:
        percentiles = [0.00, 0.01, 0.05, 0.25, 0.50, 0.75, 0.95, 0.99, 0.999, 1.00]
    rows = []
    for col in columns:
        s = pd.to_numeric(df[col], errors="coerce").replace([np.inf, -np.inf], np.nan).dropna()
        if len(s) == 0:
            continue
        q = s.quantile(percentiles)
        row = {
            "column": col,
            "count": len(s),
            "mean": s.mean(),
            "std": s.std(),
            "missing_count": df[col].isna().sum(),
            "missing_pct": df[col].isna().mean() * 100,
        }
        for p, v in q.items():
            row[f"p{p*100:g}"] = v
        rows.append(row)
    return pd.DataFrame(rows)


def psi_from_bins(expected, actual, bins):
    expected_counts = pd.cut(expected, bins=bins, include_lowest=True).value_counts(sort=False)
    actual_counts = pd.cut(actual, bins=bins, include_lowest=True).value_counts(sort=False)
    expected_pct = expected_counts / max(expected_counts.sum(), 1)
    actual_pct = actual_counts / max(actual_counts.sum(), 1)
    expected_pct = expected_pct.replace(0, 0.0001)
    actual_pct = actual_pct.replace(0, 0.0001)
    return float(((actual_pct - expected_pct) * np.log(actual_pct / expected_pct)).sum())


def numeric_psi(train_s, test_s, n_bins=10):
    train_s = pd.to_numeric(train_s, errors="coerce").replace([np.inf, -np.inf], np.nan).dropna()
    test_s = pd.to_numeric(test_s, errors="coerce").replace([np.inf, -np.inf], np.nan).dropna()
    if train_s.nunique() <= 1 or test_s.nunique() <= 1:
        return np.nan
    quantiles = np.linspace(0, 1, n_bins + 1)
    bins = np.unique(train_s.quantile(quantiles).values)
    if len(bins) < 3:
        bins = np.unique(np.linspace(train_s.min(), train_s.max(), n_bins + 1))
    if len(bins) < 3:
        return np.nan
    bins[0] = -np.inf
    bins[-1] = np.inf
    return psi_from_bins(train_s, test_s, bins)


def categorical_psi(train_s, test_s):
    train = train_s.fillna("Missing").astype(str)
    test = test_s.fillna("Missing").astype(str)
    categories = sorted(set(train.unique()) | set(test.unique()))
    train_pct = train.value_counts(normalize=True).reindex(categories, fill_value=0).replace(0, 0.0001)
    test_pct = test.value_counts(normalize=True).reindex(categories, fill_value=0).replace(0, 0.0001)
    return float(((test_pct - train_pct) * np.log(test_pct / train_pct)).sum())


def psi_level(psi):
    if pd.isna(psi):
        return "not_available"
    if psi < 0.10:
        return "low"
    if psi < 0.25:
        return "medium"
    return "high"


# ============================================================
# 1. KEY QUALITY
# ============================================================

key_rows = []
for file_name, key_cols in KEY_SPECS.items():
    key_rows.extend(key_quality_profile(file_name, key_cols))

key_quality_df = pd.DataFrame(key_rows)
save_table(key_quality_df, "01_key_quality_profile")

plt.figure(figsize=(11, 5))
plot_df = key_quality_df.sort_values("null_key_count", ascending=False)
sns.barplot(data=plot_df, y="file", x="null_key_count", hue="key", dodge=False)
plt.title("Null Values in Key Columns")
plt.xlabel("Null key count")
plt.ylabel("File")
save_fig("01_null_values_in_key_columns")


# ============================================================
# 2. MISSING VALUES ACROSS TABLES
# ============================================================

missing_tables = []
for file_name in FILES.values():
    if file_name == "HomeCredit_columns_description.csv":
        continue
    print(f"Profiling missing values: {file_name}")
    missing_tables.append(exact_missing_profile(file_name))

missing_all = pd.concat(missing_tables, ignore_index=True)
save_table(missing_all, "02_missing_values_all_columns")

missing_summary = (
    missing_all
    .assign(has_missing=lambda x: x["missing_count"] > 0)
    .groupby("file")
    .agg(
        columns=("column", "count"),
        columns_with_missing=("has_missing", "sum"),
        avg_missing_pct=("missing_pct", "mean"),
        max_missing_pct=("missing_pct", "max"),
    )
    .reset_index()
    .sort_values("max_missing_pct", ascending=False)
)
save_table(missing_summary, "03_missing_summary_by_table")

plt.figure(figsize=(12, 6))
plot_df = missing_summary.sort_values("max_missing_pct", ascending=True)
plt.barh(plot_df["file"], plot_df["max_missing_pct"])
plt.xlabel("Max missing percentage in table (%)")
plt.ylabel("File")
plt.title("Maximum Column Missing Rate by Table")
save_fig("02_max_missing_rate_by_table")

top_missing = missing_all.sort_values("missing_pct", ascending=False).head(30)
save_table(top_missing, "04_top30_missing_columns_overall")

plt.figure(figsize=(11, 8))
plot_df = top_missing.sort_values("missing_pct", ascending=True)
sns.barplot(data=plot_df, y="column", x="missing_pct", hue="file", dodge=False)
plt.xlabel("Missing percentage (%)")
plt.ylabel("Column")
plt.title("Top 30 Missing Columns Across Tables")
save_fig("03_top30_missing_columns_overall")


# ============================================================
# 3. APPLICATION TRAIN: MISSINGNESS VS TARGET
# ============================================================

app_train = pd.read_csv(DATA_PATH / "application_train.csv")
app_test = pd.read_csv(DATA_PATH / "application_test.csv")

app_missing = missing_all[missing_all["file"] == "application_train.csv"].copy()
top_app_missing_cols = app_missing[app_missing["missing_count"] > 0].head(40)["column"].tolist()

missing_target_rows = []
for col in top_app_missing_cols:
    is_missing = app_train[col].isna()
    n_missing = int(is_missing.sum())
    n_non_missing = int((~is_missing).sum())
    default_missing = app_train.loc[is_missing, "TARGET"].mean() if n_missing > 0 else np.nan
    default_non_missing = app_train.loc[~is_missing, "TARGET"].mean() if n_non_missing > 0 else np.nan
    missing_target_rows.append({
        "column": col,
        "missing_count": n_missing,
        "missing_pct": n_missing / len(app_train) * 100,
        "default_rate_when_missing": default_missing * 100 if pd.notna(default_missing) else np.nan,
        "default_rate_when_not_missing": default_non_missing * 100 if pd.notna(default_non_missing) else np.nan,
        "default_rate_diff_pp": (default_missing - default_non_missing) * 100 if pd.notna(default_missing) and pd.notna(default_non_missing) else np.nan,
    })

missing_target_df = pd.DataFrame(missing_target_rows).sort_values("missing_pct", ascending=False)
save_table(missing_target_df, "05_application_missingness_vs_target")

plt.figure(figsize=(10, 8))
plot_df = missing_target_df.dropna(subset=["default_rate_diff_pp"]).copy()
plot_df = plot_df.reindex(plot_df["default_rate_diff_pp"].abs().sort_values(ascending=False).index).head(20)
plot_df = plot_df.sort_values("default_rate_diff_pp")
sns.barplot(data=plot_df, y="column", x="default_rate_diff_pp")
plt.axvline(0, color="black", linewidth=1)
plt.xlabel("Default rate difference: missing - non-missing (percentage points)")
plt.ylabel("Column")
plt.title("Missingness Signal: Default Rate Difference")
save_fig("04_missingness_vs_target_default_rate_diff")


# ============================================================
# 4. SPECIAL VALUES IN APPLICATION
# ============================================================

special_rows = []

def add_special(metric, count, total, note):
    special_rows.append({
        "metric": metric,
        "count": int(count),
        "total_rows": int(total),
        "pct": count / total * 100,
        "note": note,
    })

add_special("CODE_GENDER_XNA", (app_train["CODE_GENDER"] == "XNA").sum(), len(app_train), "Unknown/invalid-like gender code")
add_special("DAYS_EMPLOYED_365243", (app_train["DAYS_EMPLOYED"] == 365243).sum(), len(app_train), "Special anomalous value")
add_special("AMT_ANNUITY_missing", app_train["AMT_ANNUITY"].isna().sum(), len(app_train), "Missing annuity")
add_special("ORGANIZATION_TYPE_XNA", (app_train["ORGANIZATION_TYPE"] == "XNA").sum(), len(app_train), "May mean not applicable, not necessarily invalid")
add_special("NAME_FAMILY_STATUS_Unknown", (app_train["NAME_FAMILY_STATUS"] == "Unknown").sum(), len(app_train), "Rare unknown category")

income_p99 = app_train["AMT_INCOME_TOTAL"].quantile(0.99)
income_p999 = app_train["AMT_INCOME_TOTAL"].quantile(0.999)
add_special("AMT_INCOME_TOTAL_gt_p99", (app_train["AMT_INCOME_TOTAL"] > income_p99).sum(), len(app_train), f"Income above p99={income_p99:.2f}")
add_special("AMT_INCOME_TOTAL_gt_p999", (app_train["AMT_INCOME_TOTAL"] > income_p999).sum(), len(app_train), f"Income above p99.9={income_p999:.2f}")

special_df = pd.DataFrame(special_rows).sort_values("count", ascending=False)
save_table(special_df, "06_application_special_values")

special_target_rows = []
for metric, mask in [
    ("CODE_GENDER_XNA", app_train["CODE_GENDER"] == "XNA"),
    ("DAYS_EMPLOYED_365243", app_train["DAYS_EMPLOYED"] == 365243),
    ("ORGANIZATION_TYPE_XNA", app_train["ORGANIZATION_TYPE"] == "XNA"),
    ("AMT_INCOME_TOTAL_gt_p99", app_train["AMT_INCOME_TOTAL"] > income_p99),
]:
    special_target_rows.append({
        "metric": metric,
        "count": int(mask.sum()),
        "default_rate_special": app_train.loc[mask, "TARGET"].mean() * 100 if mask.sum() > 0 else np.nan,
        "default_rate_other": app_train.loc[~mask, "TARGET"].mean() * 100 if (~mask).sum() > 0 else np.nan,
    })

special_target_df = pd.DataFrame(special_target_rows)
special_target_df["default_rate_diff_pp"] = special_target_df["default_rate_special"] - special_target_df["default_rate_other"]
save_table(special_target_df, "07_special_values_vs_target")

plt.figure(figsize=(10, 5))
plot_df = special_df[special_df["count"] > 0].sort_values("count", ascending=True)
sns.barplot(data=plot_df, y="metric", x="count")
plt.xlabel("Count")
plt.ylabel("Special value")
plt.title("Special / Invalid-like Values in application_train")
save_fig("05_application_special_values")


# ============================================================
# 5. TIME-LIKE VARIABLE QUALITY
# ============================================================

time_rows = []
for file_name in [
    "application_train.csv",
    "application_test.csv",
    "bureau.csv",
    "bureau_balance.csv",
    "previous_application.csv",
    "POS_CASH_balance.csv",
    "installments_payments.csv",
    "credit_card_balance.csv",
]:
    cols = [c for c in read_header(file_name) if ("DAYS" in c or "MONTHS" in c)]
    if not cols:
        continue
    stats = {col: {"min": np.inf, "max": -np.inf, "missing": 0, "non_missing": 0, "positive_count": 0, "zero_count": 0} for col in cols}
    for chunk in read_csv_chunks(file_name, usecols=cols, chunksize=500_000):
        for col in cols:
            s = pd.to_numeric(chunk[col], errors="coerce")
            stats[col]["missing"] += int(s.isna().sum())
            stats[col]["non_missing"] += int(s.notna().sum())
            stats[col]["positive_count"] += int((s > 0).sum())
            stats[col]["zero_count"] += int((s == 0).sum())
            if s.notna().any():
                stats[col]["min"] = min(stats[col]["min"], float(s.min()))
                stats[col]["max"] = max(stats[col]["max"], float(s.max()))
    for col, st in stats.items():
        time_rows.append({
            "file": file_name,
            "column": col,
            "min": st["min"],
            "max": st["max"],
            "missing": st["missing"],
            "non_missing": st["non_missing"],
            "positive_count": st["positive_count"],
            "positive_pct": safe_divide(st["positive_count"], st["non_missing"]) * 100 if st["non_missing"] else np.nan,
            "zero_count": st["zero_count"],
        })

time_quality_df = pd.DataFrame(time_rows).sort_values(["file", "column"])
save_table(time_quality_df, "08_time_like_variable_quality")

plt.figure(figsize=(12, 8))
plot_df = time_quality_df[time_quality_df["positive_count"] > 0].sort_values("positive_count", ascending=True)
if len(plot_df) > 0:
    sns.barplot(data=plot_df, y="column", x="positive_count", hue="file", dodge=False)
    plt.xlabel("Positive value count")
    plt.ylabel("Time-like column")
    plt.title("Positive Values in Time-like Variables")
else:
    plt.text(0.5, 0.5, "No positive time-like values found", ha="center", va="center")
    plt.axis("off")
save_fig("06_positive_time_like_values")


# ============================================================
# 6. PLAUSIBILITY CHECKS BY TABLE
# ============================================================

plausibility_rows = []

def add_plausibility(file_name, check_name, count, total, severity, note):
    plausibility_rows.append({
        "file": file_name,
        "check": check_name,
        "count": int(count),
        "total_rows": int(total),
        "pct": count / total * 100 if total else np.nan,
        "severity": severity,
        "note": note,
    })


# bureau.csv
total = 0
checks = defaultdict(int)
cols = ["CREDIT_DAY_OVERDUE", "AMT_CREDIT_SUM", "AMT_CREDIT_SUM_DEBT", "AMT_CREDIT_SUM_OVERDUE", "CREDIT_ACTIVE"]
for chunk in read_csv_chunks("bureau.csv", usecols=cols):
    total += len(chunk)
    overdue = pd.to_numeric(chunk["CREDIT_DAY_OVERDUE"], errors="coerce")
    credit_sum = pd.to_numeric(chunk["AMT_CREDIT_SUM"], errors="coerce")
    debt = pd.to_numeric(chunk["AMT_CREDIT_SUM_DEBT"], errors="coerce")
    overdue_amt = pd.to_numeric(chunk["AMT_CREDIT_SUM_OVERDUE"], errors="coerce")
    checks["negative_credit_day_overdue"] += int((overdue < 0).sum())
    checks["negative_credit_sum"] += int((credit_sum < 0).sum())
    checks["negative_debt"] += int((debt < 0).sum())
    checks["debt_greater_than_credit_sum"] += int((debt > credit_sum).sum())
    checks["negative_overdue_amount"] += int((overdue_amt < 0).sum())
    checks["bad_debt_status"] += int((chunk["CREDIT_ACTIVE"] == "Bad debt").sum())

for name, count in checks.items():
    add_plausibility("bureau.csv", name, count, total, "review", "Business plausibility check")


# previous_application.csv
total = 0
checks = defaultdict(int)
prev_cols = [
    "AMT_APPLICATION", "AMT_CREDIT", "AMT_ANNUITY",
    "DAYS_FIRST_DRAWING", "DAYS_FIRST_DUE", "DAYS_LAST_DUE_1ST_VERSION", "DAYS_LAST_DUE", "DAYS_TERMINATION",
]
prev_cols = [c for c in prev_cols if c in read_header("previous_application.csv")]
for chunk in read_csv_chunks("previous_application.csv", usecols=prev_cols):
    total += len(chunk)
    for col in [c for c in prev_cols if c.startswith("AMT_")]:
        s = pd.to_numeric(chunk[col], errors="coerce")
        checks[f"{col}_negative"] += int((s < 0).sum())
        checks[f"{col}_zero"] += int((s == 0).sum())
    for col in [c for c in prev_cols if c.startswith("DAYS_")]:
        s = pd.to_numeric(chunk[col], errors="coerce")
        checks[f"{col}_365243"] += int((s == 365243).sum())

for name, count in checks.items():
    severity = "high" if name.endswith("_365243") else "review"
    add_plausibility("previous_application.csv", name, count, total, severity, "Special date or amount plausibility check")


# POS_CASH_balance.csv
total = 0
checks = defaultdict(int)
cols = ["CNT_INSTALMENT", "CNT_INSTALMENT_FUTURE", "SK_DPD", "SK_DPD_DEF"]
for chunk in read_csv_chunks("POS_CASH_balance.csv", usecols=cols):
    total += len(chunk)
    cnt = pd.to_numeric(chunk["CNT_INSTALMENT"], errors="coerce")
    cnt_future = pd.to_numeric(chunk["CNT_INSTALMENT_FUTURE"], errors="coerce")
    dpd = pd.to_numeric(chunk["SK_DPD"], errors="coerce")
    dpd_def = pd.to_numeric(chunk["SK_DPD_DEF"], errors="coerce")
    checks["negative_cnt_instalment"] += int((cnt < 0).sum())
    checks["future_instalment_greater_than_total"] += int((cnt_future > cnt).sum())
    checks["negative_sk_dpd"] += int((dpd < 0).sum())
    checks["negative_sk_dpd_def"] += int((dpd_def < 0).sum())
    checks["sk_dpd_def_greater_than_sk_dpd"] += int((dpd_def > dpd).sum())

for name, count in checks.items():
    add_plausibility("POS_CASH_balance.csv", name, count, total, "review", "Installment count and DPD plausibility")


# installments_payments.csv
total = 0
checks = defaultdict(int)
cols = ["DAYS_INSTALMENT", "DAYS_ENTRY_PAYMENT", "AMT_INSTALMENT", "AMT_PAYMENT"]
for chunk in read_csv_chunks("installments_payments.csv", usecols=cols):
    total += len(chunk)
    days_inst = pd.to_numeric(chunk["DAYS_INSTALMENT"], errors="coerce")
    days_pay = pd.to_numeric(chunk["DAYS_ENTRY_PAYMENT"], errors="coerce")
    amt_inst = pd.to_numeric(chunk["AMT_INSTALMENT"], errors="coerce")
    amt_pay = pd.to_numeric(chunk["AMT_PAYMENT"], errors="coerce")
    checks["missing_payment_date"] += int(days_pay.isna().sum())
    checks["missing_payment_amount"] += int(amt_pay.isna().sum())
    checks["late_payment_rows"] += int((days_pay > days_inst).sum())
    checks["underpayment_rows"] += int((amt_pay < amt_inst).sum())
    checks["negative_instalment_amount"] += int((amt_inst < 0).sum())
    checks["negative_payment_amount"] += int((amt_pay < 0).sum())
    checks["zero_instalment_amount"] += int((amt_inst == 0).sum())
    checks["zero_payment_amount"] += int((amt_pay == 0).sum())

for name, count in checks.items():
    severity = "signal" if name in ["late_payment_rows", "underpayment_rows"] else "review"
    add_plausibility("installments_payments.csv", name, count, total, severity, "Payment behavior or payment data quality check")


# credit_card_balance.csv
total = 0
checks = defaultdict(int)
cols = ["AMT_BALANCE", "AMT_CREDIT_LIMIT_ACTUAL", "AMT_TOTAL_RECEIVABLE", "SK_DPD", "SK_DPD_DEF"]
for chunk in read_csv_chunks("credit_card_balance.csv", usecols=cols):
    total += len(chunk)
    bal = pd.to_numeric(chunk["AMT_BALANCE"], errors="coerce")
    limit = pd.to_numeric(chunk["AMT_CREDIT_LIMIT_ACTUAL"], errors="coerce")
    receivable = pd.to_numeric(chunk["AMT_TOTAL_RECEIVABLE"], errors="coerce")
    dpd = pd.to_numeric(chunk["SK_DPD"], errors="coerce")
    dpd_def = pd.to_numeric(chunk["SK_DPD_DEF"], errors="coerce")
    utilization = bal / limit.replace(0, np.nan)
    checks["negative_balance"] += int((bal < 0).sum())
    checks["negative_total_receivable"] += int((receivable < 0).sum())
    checks["zero_credit_limit"] += int((limit == 0).sum())
    checks["positive_balance_with_zero_limit"] += int(((bal > 0) & (limit == 0)).sum())
    checks["credit_utilization_gt_1"] += int((utilization > 1).sum())
    checks["negative_sk_dpd"] += int((dpd < 0).sum())
    checks["negative_sk_dpd_def"] += int((dpd_def < 0).sum())
    checks["sk_dpd_def_greater_than_sk_dpd"] += int((dpd_def > dpd).sum())

for name, count in checks.items():
    add_plausibility("credit_card_balance.csv", name, count, total, "review", "Credit card balance plausibility check")

plausibility_df = pd.DataFrame(plausibility_rows).sort_values(["file", "count"], ascending=[True, False])
save_table(plausibility_df, "09_plausibility_checks")

plt.figure(figsize=(12, 8))
plot_df = plausibility_df[plausibility_df["count"] > 0].sort_values("count", ascending=True).tail(25)
sns.barplot(data=plot_df, y="check", x="count", hue="file", dodge=False)
plt.xlabel("Count")
plt.ylabel("Check")
plt.title("Top Plausibility / Special-Value Checks")
save_fig("07_top_plausibility_checks")


# ============================================================
# 7. OUTLIER PROFILE FOR APPLICATION CORE VARIABLES
# ============================================================

app_train["AGE_YEARS"] = -app_train["DAYS_BIRTH"] / 365.25
app_train["DAYS_EMPLOYED_CLEAN"] = app_train["DAYS_EMPLOYED"].replace(365243, np.nan)
app_train["EMPLOYED_YEARS"] = -app_train["DAYS_EMPLOYED_CLEAN"] / 365.25
app_train["CREDIT_INCOME_RATIO"] = app_train["AMT_CREDIT"] / app_train["AMT_INCOME_TOTAL"].replace(0, np.nan)
app_train["ANNUITY_INCOME_RATIO"] = app_train["AMT_ANNUITY"] / app_train["AMT_INCOME_TOTAL"].replace(0, np.nan)

core_numeric_cols = [
    "AGE_YEARS",
    "EMPLOYED_YEARS",
    "AMT_INCOME_TOTAL",
    "AMT_CREDIT",
    "AMT_ANNUITY",
    "CREDIT_INCOME_RATIO",
    "ANNUITY_INCOME_RATIO",
]
outlier_profile_df = numeric_percentile_profile(app_train, core_numeric_cols)
save_table(outlier_profile_df, "10_application_core_outlier_profile")

fig, axes = plt.subplots(2, 2, figsize=(14, 10))
sns.histplot(app_train["AMT_INCOME_TOTAL"].clip(upper=app_train["AMT_INCOME_TOTAL"].quantile(0.99)), bins=50, ax=axes[0, 0])
axes[0, 0].set_title("Income - clipped at p99 for visualization")
sns.histplot(app_train["AMT_CREDIT"], bins=50, ax=axes[0, 1])
axes[0, 1].set_title("Credit Amount")
sns.histplot(app_train["CREDIT_INCOME_RATIO"].clip(upper=app_train["CREDIT_INCOME_RATIO"].quantile(0.99)), bins=50, ax=axes[1, 0])
axes[1, 0].set_title("Credit-Income Ratio - clipped at p99")
sns.histplot(app_train["ANNUITY_INCOME_RATIO"].clip(upper=app_train["ANNUITY_INCOME_RATIO"].quantile(0.99)), bins=50, ax=axes[1, 1])
axes[1, 1].set_title("Annuity-Income Ratio - clipped at p99")
save_fig("08_application_outlier_visual_check")


# ============================================================
# 8. TRAIN-TEST DRIFT CHECK
# ============================================================

numeric_drift_cols = [
    "CNT_CHILDREN",
    "AMT_INCOME_TOTAL",
    "AMT_CREDIT",
    "AMT_ANNUITY",
    "AMT_GOODS_PRICE",
    "REGION_POPULATION_RELATIVE",
    "DAYS_BIRTH",
    "DAYS_EMPLOYED",
    "EXT_SOURCE_1",
    "EXT_SOURCE_2",
    "EXT_SOURCE_3",
]
numeric_drift_cols = [c for c in numeric_drift_cols if c in app_train.columns and c in app_test.columns]

categorical_drift_cols = [
    "NAME_CONTRACT_TYPE",
    "CODE_GENDER",
    "FLAG_OWN_CAR",
    "FLAG_OWN_REALTY",
    "NAME_TYPE_SUITE",
    "NAME_INCOME_TYPE",
    "NAME_EDUCATION_TYPE",
    "NAME_FAMILY_STATUS",
    "NAME_HOUSING_TYPE",
    "OCCUPATION_TYPE",
    "ORGANIZATION_TYPE",
]
categorical_drift_cols = [c for c in categorical_drift_cols if c in app_train.columns and c in app_test.columns]

drift_rows = []
for col in numeric_drift_cols:
    train_s = app_train[col].replace(365243, np.nan) if col == "DAYS_EMPLOYED" else app_train[col]
    test_s = app_test[col].replace(365243, np.nan) if col == "DAYS_EMPLOYED" else app_test[col]
    psi = numeric_psi(train_s, test_s)
    drift_rows.append({
        "column": col,
        "type": "numeric",
        "train_missing_pct": app_train[col].isna().mean() * 100,
        "test_missing_pct": app_test[col].isna().mean() * 100,
        "train_mean_or_top_pct": pd.to_numeric(train_s, errors="coerce").mean(),
        "test_mean_or_top_pct": pd.to_numeric(test_s, errors="coerce").mean(),
        "psi": psi,
        "psi_level": psi_level(psi),
    })

for col in categorical_drift_cols:
    psi = categorical_psi(app_train[col], app_test[col])
    train_top_pct = app_train[col].fillna("Missing").astype(str).value_counts(normalize=True).iloc[0] * 100
    test_top_pct = app_test[col].fillna("Missing").astype(str).value_counts(normalize=True).iloc[0] * 100
    drift_rows.append({
        "column": col,
        "type": "categorical",
        "train_missing_pct": app_train[col].isna().mean() * 100,
        "test_missing_pct": app_test[col].isna().mean() * 100,
        "train_mean_or_top_pct": train_top_pct,
        "test_mean_or_top_pct": test_top_pct,
        "psi": psi,
        "psi_level": psi_level(psi),
    })

drift_df = pd.DataFrame(drift_rows).sort_values("psi", ascending=False)
save_table(drift_df, "11_train_test_drift_psi")

plt.figure(figsize=(12, 7))
plot_df = drift_df.dropna(subset=["psi"]).sort_values("psi", ascending=True)
sns.barplot(data=plot_df, y="column", x="psi", hue="psi_level", dodge=False)
plt.axvline(0.10, color="orange", linestyle="--", linewidth=1, label="medium threshold")
plt.axvline(0.25, color="red", linestyle="--", linewidth=1, label="high threshold")
plt.xlabel("PSI")
plt.ylabel("Column")
plt.title("Train-Test Drift Check Using PSI")
plt.legend()
save_fig("09_train_test_drift_psi")


# ============================================================
# 9. CLEANING RULES
# ============================================================

cleaning_rules = pd.DataFrame([
    {
        "issue": "DAYS_EMPLOYED equals 365243",
        "affected_tables": "application_train/test",
        "recommended_action": "Replace 365243 with NaN and add DAYS_EMPLOYED_ANOM indicator",
        "why": "365243 is inconsistent with negative day variables and likely represents a special code",
        "alternative": "Drop rows or keep raw",
        "pros_cons": "Indicator keeps anomaly signal; dropping rows loses many records; keeping raw distorts tenure analysis",
    },
    {
        "issue": "Previous application date columns equal 365243",
        "affected_tables": "previous_application",
        "recommended_action": "Replace 365243 with NaN in DAYS_FIRST_DRAWING, DAYS_FIRST_DUE, DAYS_LAST_DUE_1ST_VERSION, DAYS_LAST_DUE, DAYS_TERMINATION and add indicators if used",
        "why": "365243 appears as a special date value and should not be treated as real elapsed days",
        "alternative": "Keep raw date values",
        "pros_cons": "Replacing prevents distorted time features; indicators preserve the special-value signal; keeping raw corrupts averages and ranges",
    },
    {
        "issue": "CODE_GENDER equals XNA",
        "affected_tables": "application_train",
        "recommended_action": "Treat as Unknown/NaN; exclude from gender-specific EDA if needed",
        "why": "Only a few records and not a standard gender category",
        "alternative": "Drop rows or map to mode",
        "pros_cons": "Unknown is transparent; mode hides data issue; dropping is acceptable only because count is tiny",
    },
    {
        "issue": "High missing building/apartment variables",
        "affected_tables": "application_train/test",
        "recommended_action": "Do not drop automatically; analyze missingness signal and consider missing indicators",
        "why": "Missingness may reflect unavailable property information and can carry risk signal",
        "alternative": "Drop high-missing columns",
        "pros_cons": "Dropping simplifies model but may remove useful missingness signal",
    },
    {
        "issue": "Property/building missingness is related to TARGET",
        "affected_tables": "application_train/test",
        "recommended_action": "Create property information missing indicators or a property_missing_count feature",
        "why": "Several property variables show higher default rate when missing, so missingness itself can be informative",
        "alternative": "Drop property columns or only impute numeric values",
        "pros_cons": "Missing indicators preserve risk signal; dropping loses signal; blind imputation hides data availability differences",
    },
    {
        "issue": "OWN_CAR_AGE missing",
        "affected_tables": "application_train/test",
        "recommended_action": "Interpret together with FLAG_OWN_CAR; missing likely means no car for many customers",
        "why": "Missing is not always data error; it can be not applicable",
        "alternative": "Median imputation",
        "pros_cons": "Business-aware handling avoids imputing car age for customers without cars",
    },
    {
        "issue": "Monetary outliers",
        "affected_tables": "application, bureau, previous, installments, credit card",
        "recommended_action": "Keep raw for analysis/modeling, use clipping/log scale for visualization",
        "why": "Extreme financial values can be real and risk-relevant",
        "alternative": "Winsorize or remove outliers",
        "pros_cons": "Keeping raw preserves signal; clipping improves chart readability; removal can bias analysis",
    },
    {
        "issue": "Credit card credit limit equals zero",
        "affected_tables": "credit_card_balance",
        "recommended_action": "When computing utilization, avoid direct division by zero and create zero_credit_limit_flag",
        "why": "Zero credit limit can create infinite/undefined utilization and may itself be an informative credit card state",
        "alternative": "Drop zero-limit rows or replace limit with a small constant",
        "pros_cons": "Flagging keeps the signal and avoids invalid ratios; dropping loses history; small constants create artificial extreme utilization",
    },
    {
        "issue": "Historical tables have one-to-many records",
        "affected_tables": "bureau, bureau_balance, previous_application, POS, installments, credit card",
        "recommended_action": "Do not drop repeated SK_ID_CURR; aggregate to SK_ID_CURR in Step 4",
        "why": "Repeated records are valid history, not duplicate error",
        "alternative": "Raw merge",
        "pros_cons": "Aggregation preserves customer grain; raw merge duplicates customers and biases metrics",
    },
    {
        "issue": "Missing after merging historical features",
        "affected_tables": "final analytical table",
        "recommended_action": "Use history indicator and fill count-like aggregates with 0 where business-valid",
        "why": "Missing may mean no history in that source",
        "alternative": "Generic mean/median imputation",
        "pros_cons": "Business-aware fill is interpretable; generic imputation can blur no-history customers",
    },
    {
        "issue": "Imbalanced target",
        "affected_tables": "application_train",
        "recommended_action": "Use default rate in EDA and ROC-AUC/Recall/Precision/F1 in ML",
        "why": "TARGET=1 is minority class",
        "alternative": "Accuracy only",
        "pros_cons": "Accuracy can be misleading in imbalanced credit risk data",
    },
    {
        "issue": "Rare categorical values",
        "affected_tables": "all categorical tables",
        "recommended_action": "Show sample size with default rate; group rare categories for visualization if needed",
        "why": "Rare categories can show unstable default rates",
        "alternative": "One-hot every category",
        "pros_cons": "Rare grouping improves stability; full one-hot preserves detail but can overfit",
    },
    {
        "issue": "Potential time leakage",
        "affected_tables": "all historical tables",
        "recommended_action": "Use only information available at or before current application time",
        "why": "Future information creates unrealistic analysis/model performance",
        "alternative": "Use all historical columns blindly",
        "pros_cons": "Leakage guardrail keeps project defensible; blind usage risks inflated conclusions",
    },
])

save_table(cleaning_rules, "12_cleaning_rules")


# ============================================================
# 10. CLEAN APPLICATION BASE PREVIEW
# ============================================================

def clean_application_base(df):
    df = df.copy()
    df["DAYS_EMPLOYED_ANOM"] = (df["DAYS_EMPLOYED"] == 365243).astype(int)
    df["DAYS_EMPLOYED"] = df["DAYS_EMPLOYED"].replace(365243, np.nan)
    df["CODE_GENDER"] = df["CODE_GENDER"].replace("XNA", np.nan)

    df["AGE_YEARS"] = -df["DAYS_BIRTH"] / 365.25
    df["EMPLOYED_YEARS"] = -df["DAYS_EMPLOYED"] / 365.25
    df["REGISTRATION_YEARS"] = -df["DAYS_REGISTRATION"] / 365.25
    df["ID_PUBLISH_YEARS"] = -df["DAYS_ID_PUBLISH"] / 365.25

    df["CREDIT_INCOME_RATIO"] = df["AMT_CREDIT"] / df["AMT_INCOME_TOTAL"].replace(0, np.nan)
    df["ANNUITY_INCOME_RATIO"] = df["AMT_ANNUITY"] / df["AMT_INCOME_TOTAL"].replace(0, np.nan)
    df["GOODS_CREDIT_RATIO"] = df["AMT_GOODS_PRICE"] / df["AMT_CREDIT"].replace(0, np.nan)
    return df


clean_train_preview = clean_application_base(app_train.head(1000))
clean_preview_cols = [
    "SK_ID_CURR",
    "TARGET",
    "CODE_GENDER",
    "DAYS_EMPLOYED_ANOM",
    "AGE_YEARS",
    "EMPLOYED_YEARS",
    "CREDIT_INCOME_RATIO",
    "ANNUITY_INCOME_RATIO",
    "GOODS_CREDIT_RATIO",
]
save_table(clean_train_preview[clean_preview_cols].head(30), "13_clean_application_base_preview")

cleaning_impact = pd.DataFrame([
    {
        "rule": "DAYS_EMPLOYED_365243_to_NaN",
        "affected_rows_train": int((app_train["DAYS_EMPLOYED"] == 365243).sum()),
        "affected_pct_train": (app_train["DAYS_EMPLOYED"] == 365243).mean() * 100,
    },
    {
        "rule": "CODE_GENDER_XNA_to_NaN",
        "affected_rows_train": int((app_train["CODE_GENDER"] == "XNA").sum()),
        "affected_pct_train": (app_train["CODE_GENDER"] == "XNA").mean() * 100,
    },
    {
        "rule": "AMT_ANNUITY_missing_kept_for_imputation",
        "affected_rows_train": int(app_train["AMT_ANNUITY"].isna().sum()),
        "affected_pct_train": app_train["AMT_ANNUITY"].isna().mean() * 100,
    },
    {
        "rule": "PROPERTY_INFO_MISSING_SIGNAL_keep_indicator",
        "affected_rows_train": int(app_train[["ENTRANCES_AVG", "ELEVATORS_AVG", "APARTMENTS_AVG", "LIVINGAREA_AVG", "WALLSMATERIAL_MODE"]].isna().any(axis=1).sum()),
        "affected_pct_train": app_train[["ENTRANCES_AVG", "ELEVATORS_AVG", "APARTMENTS_AVG", "LIVINGAREA_AVG", "WALLSMATERIAL_MODE"]].isna().any(axis=1).mean() * 100,
    },
    {
        "rule": "PREVIOUS_APPLICATION_DAYS_FIRST_DRAWING_365243_to_NaN",
        "affected_rows_train": int(plausibility_df.loc[plausibility_df["check"] == "DAYS_FIRST_DRAWING_365243", "count"].iloc[0]) if (plausibility_df["check"] == "DAYS_FIRST_DRAWING_365243").any() else 0,
        "affected_pct_train": float(plausibility_df.loc[plausibility_df["check"] == "DAYS_FIRST_DRAWING_365243", "pct"].iloc[0]) if (plausibility_df["check"] == "DAYS_FIRST_DRAWING_365243").any() else 0,
    },
    {
        "rule": "PREVIOUS_APPLICATION_DAYS_FIRST_DUE_365243_to_NaN",
        "affected_rows_train": int(plausibility_df.loc[plausibility_df["check"] == "DAYS_FIRST_DUE_365243", "count"].iloc[0]) if (plausibility_df["check"] == "DAYS_FIRST_DUE_365243").any() else 0,
        "affected_pct_train": float(plausibility_df.loc[plausibility_df["check"] == "DAYS_FIRST_DUE_365243", "pct"].iloc[0]) if (plausibility_df["check"] == "DAYS_FIRST_DUE_365243").any() else 0,
    },
    {
        "rule": "PREVIOUS_APPLICATION_DAYS_LAST_DUE_1ST_VERSION_365243_to_NaN",
        "affected_rows_train": int(plausibility_df.loc[plausibility_df["check"] == "DAYS_LAST_DUE_1ST_VERSION_365243", "count"].iloc[0]) if (plausibility_df["check"] == "DAYS_LAST_DUE_1ST_VERSION_365243").any() else 0,
        "affected_pct_train": float(plausibility_df.loc[plausibility_df["check"] == "DAYS_LAST_DUE_1ST_VERSION_365243", "pct"].iloc[0]) if (plausibility_df["check"] == "DAYS_LAST_DUE_1ST_VERSION_365243").any() else 0,
    },
    {
        "rule": "PREVIOUS_APPLICATION_DAYS_LAST_DUE_365243_to_NaN",
        "affected_rows_train": int(plausibility_df.loc[plausibility_df["check"] == "DAYS_LAST_DUE_365243", "count"].iloc[0]) if (plausibility_df["check"] == "DAYS_LAST_DUE_365243").any() else 0,
        "affected_pct_train": float(plausibility_df.loc[plausibility_df["check"] == "DAYS_LAST_DUE_365243", "pct"].iloc[0]) if (plausibility_df["check"] == "DAYS_LAST_DUE_365243").any() else 0,
    },
    {
        "rule": "PREVIOUS_APPLICATION_DAYS_TERMINATION_365243_to_NaN",
        "affected_rows_train": int(plausibility_df.loc[plausibility_df["check"] == "DAYS_TERMINATION_365243", "count"].iloc[0]) if (plausibility_df["check"] == "DAYS_TERMINATION_365243").any() else 0,
        "affected_pct_train": float(plausibility_df.loc[plausibility_df["check"] == "DAYS_TERMINATION_365243", "pct"].iloc[0]) if (plausibility_df["check"] == "DAYS_TERMINATION_365243").any() else 0,
    },
    {
        "rule": "CREDIT_CARD_ZERO_LIMIT_flag_and_safe_utilization",
        "affected_rows_train": int(plausibility_df.loc[plausibility_df["check"] == "zero_credit_limit", "count"].iloc[0]) if (plausibility_df["check"] == "zero_credit_limit").any() else 0,
        "affected_pct_train": float(plausibility_df.loc[plausibility_df["check"] == "zero_credit_limit", "pct"].iloc[0]) if (plausibility_df["check"] == "zero_credit_limit").any() else 0,
    },
])
save_table(cleaning_impact, "14_cleaning_impact_summary")


# ============================================================
# 11. WORD / HTML REPORT
# ============================================================

report_path = build_step3_word_report()


# ============================================================
# 12. WRITE SUMMARY + ZIP
# ============================================================

summary = {
    "step": "Step 3 - Data Quality Assessment & Cleaning Strategy",
    "output_path": str(OUT_PATH),
    "tables_path": str(TABLE_PATH),
    "figures_path": str(FIG_PATH),
    "report_path": str(report_path),
    "n_tables_saved": len(list(TABLE_PATH.glob("*.csv"))),
    "n_figures_saved": len(list(FIG_PATH.glob("*.png"))),
    "important_note": "This step defines and audits cleaning rules. Full historical aggregation is Step 4.",
}

with open(OUT_PATH / "step3_summary.json", "w", encoding="utf-8") as f:
    json.dump(summary, f, ensure_ascii=False, indent=2)

zip_path = DATA_PATH / "step3_outputs.zip"
with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
    for file in OUT_PATH.rglob("*"):
        z.write(file, file.relative_to(OUT_PATH.parent))

print("=" * 90)
print("STEP 3 RESULT-ONLY RUN FINISHED")
print(f"Tables saved: {TABLE_PATH}")
print(f"Figures saved: {FIG_PATH}")
print(f"Zip file: {zip_path}")
print("=" * 90)
